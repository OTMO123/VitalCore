"""
Text Extraction Service

Extracts text from various document formats including PDFs, Word documents,
and other text-based formats. Complements OCR service for non-image documents.
"""

import asyncio
import io
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Optional, Dict, Any, List
from pathlib import Path
import mimetypes

import structlog
import fitz  # PyMuPDF
from xml.etree import ElementTree as ET

logger = structlog.get_logger(__name__)


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    
    text: str
    format: str
    metadata: Dict[str, Any]
    extraction_method: str
    success: bool
    error: Optional[str] = None


class TextExtractorInterface(ABC):
    """Interface for text extractors following SOLID principles."""
    
    @abstractmethod
    async def extract_text(self, file_data: bytes, mime_type: str) -> ExtractionResult:
        """Extract text from file data."""
        pass
    
    @abstractmethod
    def supports_format(self, mime_type: str) -> bool:
        """Check if format is supported."""
        pass


class PDFTextExtractor(TextExtractorInterface):
    """Extract text from PDF documents."""
    
    def __init__(self):
        self.logger = logger.bind(extractor="PDFTextExtractor")
    
    async def extract_text(self, file_data: bytes, mime_type: str) -> ExtractionResult:
        """Extract text from PDF using PyMuPDF."""
        try:
            # Open PDF document
            pdf_document = fitz.open(stream=file_data, filetype="pdf")
            
            text_content = []
            metadata = {
                "page_count": len(pdf_document),
                "title": pdf_document.metadata.get("title", ""),
                "author": pdf_document.metadata.get("author", ""),
                "subject": pdf_document.metadata.get("subject", ""),
                "creator": pdf_document.metadata.get("creator", ""),
                "producer": pdf_document.metadata.get("producer", ""),
                "creation_date": pdf_document.metadata.get("creationDate", ""),
                "modification_date": pdf_document.metadata.get("modDate", ""),
                "encrypted": pdf_document.needs_pass,
                "has_forms": pdf_document.is_form_pdf
            }
            
            # Extract text from each page
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            pdf_document.close()
            
            full_text = "\n\n".join(text_content)
            
            return ExtractionResult(
                text=full_text,
                format="pdf",
                metadata=metadata,
                extraction_method="pymupdf",
                success=True
            )
            
        except Exception as e:
            self.logger.error("PDF text extraction failed", error=str(e))
            return ExtractionResult(
                text="",
                format="pdf",
                metadata={"error": str(e)},
                extraction_method="pymupdf",
                success=False,
                error=str(e)
            )
    
    def supports_format(self, mime_type: str) -> bool:
        """Check if PDF format is supported."""
        return mime_type == "application/pdf"


class WordTextExtractor(TextExtractorInterface):
    """Extract text from Word documents."""
    
    def __init__(self):
        self.logger = logger.bind(extractor="WordTextExtractor")
    
    async def extract_text(self, file_data: bytes, mime_type: str) -> ExtractionResult:
        """Extract text from Word documents."""
        try:
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # .docx format
                return await self._extract_docx_text(file_data)
            elif mime_type == "application/msword":
                # .doc format - more complex, requires different approach
                return await self._extract_doc_text(file_data)
            else:
                raise ValueError(f"Unsupported Word format: {mime_type}")
                
        except Exception as e:
            self.logger.error("Word text extraction failed", error=str(e))
            return ExtractionResult(
                text="",
                format="word",
                metadata={"error": str(e)},
                extraction_method="python-docx",
                success=False,
                error=str(e)
            )
    
    async def _extract_docx_text(self, file_data: bytes) -> ExtractionResult:
        """Extract text from .docx files."""
        try:
            # Try to use python-docx if available
            try:
                from docx import Document
                
                # Create document from bytes
                doc = Document(io.BytesIO(file_data))
                
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text)
                
                # Extract table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text)
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                full_text = "\n".join(text_content)
                
                return ExtractionResult(
                    text=full_text,
                    format="docx",
                    metadata={
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "extraction_method": "python-docx"
                    },
                    extraction_method="python-docx",
                    success=True
                )
                
            except ImportError:
                # Fallback: Extract from XML structure directly
                return await self._extract_docx_xml(file_data)
                
        except Exception as e:
            raise e
    
    async def _extract_docx_xml(self, file_data: bytes) -> ExtractionResult:
        """Fallback: Extract text from DOCX XML structure."""
        try:
            import zipfile
            
            # DOCX is a zip file with XML content
            with zipfile.ZipFile(io.BytesIO(file_data), 'r') as zip_file:
                # Read document.xml
                xml_content = zip_file.read('word/document.xml')
                
                # Parse XML
                root = ET.fromstring(xml_content)
                
                # Extract text nodes
                text_content = []
                for elem in root.iter():
                    if elem.text:
                        text_content.append(elem.text)
                
                full_text = " ".join(text_content)
                
                return ExtractionResult(
                    text=full_text,
                    format="docx",
                    metadata={
                        "extraction_method": "xml_parsing",
                        "fallback": True
                    },
                    extraction_method="xml_parsing",
                    success=True
                )
                
        except Exception as e:
            raise e
    
    async def _extract_doc_text(self, file_data: bytes) -> ExtractionResult:
        """Extract text from .doc files (fallback implementation)."""
        # .doc format is complex binary format
        # In production, consider using external tools like antiword or LibreOffice
        
        return ExtractionResult(
            text="[DOC format not fully supported in this implementation]",
            format="doc",
            metadata={
                "extraction_method": "fallback",
                "note": "Use LibreOffice or antiword for full .doc support"
            },
            extraction_method="fallback",
            success=False,
            error="DOC format requires external conversion tools"
        )
    
    def supports_format(self, mime_type: str) -> bool:
        """Check if Word format is supported."""
        return mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]


class PlainTextExtractor(TextExtractorInterface):
    """Extract text from plain text files."""
    
    def __init__(self):
        self.logger = logger.bind(extractor="PlainTextExtractor")
    
    async def extract_text(self, file_data: bytes, mime_type: str) -> ExtractionResult:
        """Extract text from plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    
                    return ExtractionResult(
                        text=text,
                        format="text",
                        metadata={
                            "encoding": encoding,
                            "character_count": len(text),
                            "line_count": text.count('\n') + 1
                        },
                        extraction_method="direct_decode",
                        success=True
                    )
                    
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail
            return ExtractionResult(
                text="",
                format="text",
                metadata={"error": "Unable to decode text with any encoding"},
                extraction_method="direct_decode",
                success=False,
                error="Encoding detection failed"
            )
            
        except Exception as e:
            self.logger.error("Text extraction failed", error=str(e))
            return ExtractionResult(
                text="",
                format="text",
                metadata={"error": str(e)},
                extraction_method="direct_decode",
                success=False,
                error=str(e)
            )
    
    def supports_format(self, mime_type: str) -> bool:
        """Check if plain text format is supported."""
        return mime_type in [
            "text/plain",
            "text/html",
            "text/xml",
            "application/xml",
            "text/csv",
            "application/json"
        ]


class TextExtractorService:
    """Main text extraction service."""
    
    def __init__(self, extractors: Optional[List[TextExtractorInterface]] = None):
        self.extractors = extractors or [
            PDFTextExtractor(),
            WordTextExtractor(),
            PlainTextExtractor()
        ]
        self.logger = logger.bind(service="TextExtractorService")
    
    async def extract_text(
        self, 
        file_data: bytes, 
        mime_type: str, 
        filename: Optional[str] = None
    ) -> ExtractionResult:
        """Extract text from file data."""
        try:
            self.logger.info(
                "Starting text extraction",
                mime_type=mime_type,
                filename=filename,
                file_size=len(file_data)
            )
            
            # Find appropriate extractor
            extractor = self._get_extractor_for_format(mime_type)
            
            if not extractor:
                return ExtractionResult(
                    text="",
                    format="unknown",
                    metadata={"mime_type": mime_type},
                    extraction_method="none",
                    success=False,
                    error=f"No extractor available for format: {mime_type}"
                )
            
            # Perform extraction
            result = await extractor.extract_text(file_data, mime_type)
            
            self.logger.info(
                "Text extraction completed",
                success=result.success,
                text_length=len(result.text),
                extraction_method=result.extraction_method
            )
            
            return result
            
        except Exception as e:
            self.logger.error("Text extraction service failed", error=str(e))
            return ExtractionResult(
                text="",
                format="error",
                metadata={"service_error": str(e)},
                extraction_method="none",
                success=False,
                error=str(e)
            )
    
    def _get_extractor_for_format(self, mime_type: str) -> Optional[TextExtractorInterface]:
        """Get appropriate extractor for file format."""
        for extractor in self.extractors:
            if extractor.supports_format(mime_type):
                return extractor
        return None
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported formats."""
        formats = []
        for extractor in self.extractors:
            # This would need to be implemented in each extractor
            # For now, return common formats
            pass
        
        return [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain",
            "text/html",
            "text/xml",
            "application/xml",
            "text/csv",
            "application/json"
        ]


# Factory function
def get_text_extractor_service(
    extractors: Optional[List[TextExtractorInterface]] = None
) -> TextExtractorService:
    """Factory function to create text extractor service."""
    return TextExtractorService(extractors)